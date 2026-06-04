from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Helper: add coloured heading ──────────────────────────────
def add_heading(text, level=1, color=RGBColor(0x1A, 0x23, 0x7E)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 11)
    return p

def add_step_heading(step_no, title):
    p = doc.add_paragraph()
    run_no = p.add_run(f"Step {step_no}:  ")
    run_no.bold = True
    run_no.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
    run_no.font.size = Pt(12)
    run_title = p.add_run(title)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    run_title.font.size = Pt(12)
    return p

def add_body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11)
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text).font.size = Pt(11)
    return p

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run("📝  Note: " + text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

def img_placeholder(label):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ Screenshot: {label} ]")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F0F0F0')
    p._p.get_or_add_pPr().append(shading)
    return p

# ═══════════════════════════════════════════════════════════════
#  COVER
# ═══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MEIL Neyveli Energy Private Limited")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Gate Pass Management System")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("NRGP – Non-Returnable Gate Pass")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Step-by-Step Process Guide")
run.font.size = Pt(14)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Version 1.0  |  Date: 03-06-2026")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  OVERVIEW
# ═══════════════════════════════════════════════════════════════
add_heading("1.  Overview", 1)
add_body(
    "A Non-Returnable Gate Pass (NRGP) is raised when materials are sent out of "
    "MEIL premises permanently (i.e., they will NOT be returned). The process "
    "involves three roles: the Requestor (User), the HOD Approver, and the Store "
    "User who generates the physical gate pass."
)
doc.add_paragraph()

# Role table
add_heading("Roles & Responsibilities", 2)
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
hdr = tbl.rows[0].cells
hdr[0].text = "Role"
hdr[1].text = "SAP User ID (example)"
hdr[2].text = "Responsibility"
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    cell._tc.get_or_add_tcPr().append(OxmlElement('w:shd'))
    shd = cell._tc.tcPr.find(qn('w:shd'))
    shd.set(qn('w:fill'), '1A237E')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')

rows_data = [
    ("Requestor (User)",  "INTRAEDGEU7",  "Creates the NRGP request and resubmits if Amendment is required"),
    ("HOD",               "INTRAEDGEU11", "Reviews and Approves / Rejects / Sends for Amendment via SAP Inbox"),
    ("Store User",        "INTRAEDGEU12", "Validates approval, generates Gate Pass, fills logistics, prints"),
]
for role, uid, resp in rows_data:
    row = tbl.add_row().cells
    row[0].text = role
    row[1].text = uid
    row[2].text = resp

doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  FLOW DIAGRAM (text)
# ═══════════════════════════════════════════════════════════════
add_heading("2.  Process Flow Summary", 1)
flow_steps = [
    "User logs in  →",
    "Opens NRGP REQUEST from Dashboard  →",
    "Fills Gate Pass Request (Plant, Vendor, Items)  →",
    "Submits Request  →",
    "HOD receives task in SAP Inbox  →",
    "HOD Approves (or sends Amendment / Rejects)  →",
    "Store User receives task in SAP Inbox  →",
    "Store User Approves  →",
    "Store User opens Out Gate Pass Generation  →",
    "Enters Request No. → status shows APPROVED  →",
    "Clicks GENERATE GATE PASS  →",
    "Fills Logistics details  →",
    "Saves & Prints Gate Pass / DC",
]
p = doc.add_paragraph()
for i, step in enumerate(flow_steps):
    run = p.add_run(step + ("  " if i < len(flow_steps)-1 else ""))
    run.font.size = Pt(10)
    if "→" in step:
        run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)

doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  STEP-BY-STEP
# ═══════════════════════════════════════════════════════════════
add_heading("3.  Detailed Step-by-Step Process", 1)
doc.add_paragraph()

# ── STEP 1 ────────────────────────────────────────────────────
add_step_heading(1, "Login to Gate Pass Management System (Requestor)")
add_body("The Requestor opens the application URL and signs in using SAP credentials.")
add_bullet("URL: http://localhost:8080/index.html")
add_bullet("Enter Username (e.g., INTRAEDGEU7) and Password")
add_bullet("Click Sign In — the system authenticates via SAP backend")
img_placeholder("Login Screen — User INTRAEDGEU7")
doc.add_paragraph()

# ── STEP 2 ────────────────────────────────────────────────────
add_step_heading(2, "Dashboard — Home Screen")
add_body(
    "After login, the Dashboard displays all available modules as tiles. "
    "The Requestor can see NRGP REQUEST, RGP REQUEST, Gate Pass Requests List, etc. "
    "The 'Gate Pass Requests List' tile shows the current count of pending requests."
)
add_bullet("Click the NRGP REQUEST tile to create a new Non-Returnable Gate Pass")
img_placeholder("Dashboard / Home Screen")
doc.add_paragraph()

# ── STEP 3 ────────────────────────────────────────────────────
add_step_heading(3, "Fill NRGP Request — Header Details")
add_body(
    "The NRGP Creation screen opens. The header section is auto-filled based on the "
    "logged-in user's profile."
)
add_bullet("Plant — auto-filled based on user login (e.g., 2301 – NEPL HEAD SITE)")
add_bullet("Company Code — auto-filled (e.g., 2300)")
add_bullet("GP Date — today's date (auto-filled, read-only)")
add_bullet("GP Type — pre-set to Non-Returnable Gatepass (NRGP)")
add_bullet("Department — auto-filled based on user profile (e.g., Civil)")
add_bullet("Vendor — search and select from the dropdown")
add_bullet("Vendor Address & Vendor GST — auto-filled on vendor selection")
img_placeholder("NRGP Creation Form — Header Details & Vendor Information")
doc.add_paragraph()

# ── STEP 4 ────────────────────────────────────────────────────
add_step_heading(4, "Fill Table of Items & Remarks")
add_body(
    "Scroll down to add the items to be dispatched. Click '+ Add Row' to add each item."
)
add_bullet("Material — enter material number or use F4 value help")
add_bullet("Material Description — auto-filled or enter manually")
add_bullet("HSN Code — mandatory field; enter HSN code or search by description")
add_bullet("Quantity — enter the quantity to be dispatched")
add_bullet("UOM — Unit of Measure (e.g., EA, NOS, KG)")
add_bullet("Rate (Rs.) — enter unit rate; Amount is auto-calculated")
add_bullet("Remarks — optional; enter up to 250 characters")
add_note("HSN Code is mandatory for all requests. Submission will be blocked if any row has an empty HSN Code.")
img_placeholder("NRGP Creation Form — Table of Items & Remarks Section")
doc.add_paragraph()

# ── STEP 5 ────────────────────────────────────────────────────
add_step_heading(5, "Submit Gate Pass Request")
add_body(
    "Once all details are filled, click the 'Submit Gate Pass Request' button at the bottom of the page."
)
add_bullet("The system validates all mandatory fields")
add_bullet("On success, a confirmation dialog appears with the generated Request Number")
add_bullet("Example: GatePass Request NUMBER 2301265035 Created successfully!")
add_bullet("The user can copy the request number using 'Copy Number' button")
img_placeholder("Success Dialog — Request 2301265035 Created Successfully")
doc.add_paragraph()

# ── STEP 6 ────────────────────────────────────────────────────
add_step_heading(6, "View Request Status in Gate Pass Request List")
add_body(
    "The Requestor can track the status of submitted requests from the 'Gate Pass Requests List' "
    "tile on the Dashboard or via the side navigation."
)
add_bullet("Status: Pending — waiting for HOD/Store approval")
add_bullet("Status: Approved — approved by both HOD and Store")
add_bullet("Status: Rejected — rejected by HOD or Store")
add_bullet("Status: Amendment — sent back for correction by HOD")
add_bullet("Use the search box to find a specific request by number")
img_placeholder("Dashboard — Gate Pass Requests List tile highlighted (58 Pending)")
img_placeholder("Gate Pass Request List — Showing Pending and Approved requests")
doc.add_paragraph()

# ── STEP 7 ────────────────────────────────────────────────────
add_step_heading(7, "HOD Approval — Login as HOD")
add_body(
    "The HOD (Head of Department) logs in separately. The HOD dashboard shows only "
    "list/approval screens — creation screens are not visible for the HOD role."
)
add_bullet("Login with HOD credentials (e.g., INTRAEDGEU11)")
add_bullet("Dashboard shows: Gate Pass Requests List, Gate Pass List, Ash GP List, Scrap Request List")
add_bullet("Pending count is visible on the Gate Pass Requests List tile")
img_placeholder("Login Screen — HOD User INTRAEDGEU11")
img_placeholder("HOD Dashboard — Restricted view (list tiles only)")
doc.add_paragraph()

# ── STEP 8 ────────────────────────────────────────────────────
add_step_heading(8, "HOD — Review & Approve via Gate Pass Request List")
add_body(
    "The HOD can approve requests directly from the Gate Pass Request List screen "
    "or via SAP My Inbox workflow."
)
add_bullet("Open Gate Pass Requests List from Dashboard")
add_bullet("Pending requests show an 'Approve' button in the last column")
add_bullet("Click Approve to approve directly from the list")
add_bullet("Alternatively, open SAP Fiori Launchpad → My Inbox for workflow-based approval")
img_placeholder("Gate Pass Request List — HOD view with Approve buttons on Pending rows")
doc.add_paragraph()

# ── STEP 9 ────────────────────────────────────────────────────
add_step_heading(9, "HOD — SAP Inbox Approval")
add_body(
    "For detailed review, the HOD uses SAP My Inbox. The gate pass request appears as "
    "a workflow task with full item details."
)
add_bullet("Open SAP Fiori Launchpad → click 'My Inbox'")
add_bullet("Find the task: 'Gate Pass Decision <Request No.>'")
add_bullet("Review: Department, Supplier, Vendor Name, Gate Pass Type, Items table")
add_bullet("Available actions: Approve / Reject / AMENDMENT / Forward / Suspend")
add_bullet("Click Approve to forward to Store User for the next level of approval")
add_note("If the HOD clicks AMENDMENT, the request is sent back to the Requestor for correction. The Requestor must edit and re-submit.")
img_placeholder("SAP Fiori Launchpad — My Inbox tile")
img_placeholder("SAP My Inbox — Gate Pass Decision 2301265035 with Approve/Reject/Amendment actions")
doc.add_paragraph()

# ── STEP 10 ────────────────────────────────────────────────────
add_step_heading(10, "Store User Approval — Login as Store User")
add_body(
    "After HOD approval, the request moves to the Store User for second-level approval. "
    "The Store User follows the same SAP Inbox approval process as the HOD."
)
add_bullet("Login with Store User credentials (e.g., INTRAEDGEU12)")
add_bullet("Open SAP My Inbox → find the Gate Pass Decision task")
add_bullet("Review the full request details")
add_bullet("Click Approve to complete the approval workflow")
add_bullet("The request status changes to APPROVED in the Gate Pass Management System")
img_placeholder("Login Screen — Store User INTRAEDGEU12")
img_placeholder("Gate Pass Request List (Store User) — Request 2301265035 Status: Approved")
doc.add_paragraph()

# ── STEP 11 ────────────────────────────────────────────────────
add_step_heading(11, "Out Gate Pass Generation — Enter Request Number")
add_body(
    "Once the request is APPROVED, the Store User opens the Out Gate Pass Generation screen "
    "to generate the physical gate pass."
)
add_bullet("Navigate to Out Gate Pass Generation from the side navigation or Dashboard")
add_bullet("Enter the Request Number (e.g., 2301265035) in the search field")
add_bullet("Click the Search icon — system fetches request details from the backend")
add_bullet("Approval Status shows: ✅ APPROVED")
add_bullet("Header Details are auto-loaded: Plant, Company Code, Department, Vendor, Vendor Address, Vendor GST")
img_placeholder("Out Gate Pass Generation — Request Validation APPROVED + Header Details loaded")
doc.add_paragraph()

# ── STEP 12 ────────────────────────────────────────────────────
add_step_heading(12, "Generate Gate Pass")
add_body(
    "After validating the request, click the GENERATE GATE PASS button at the bottom of the screen."
)
add_bullet("Button is visible only for Store Users when the request status is APPROVED")
add_bullet("System creates the Gate Pass in the backend")
add_bullet("A success dialog displays the generated Gate Pass Number")
add_bullet("Example: Gate Pass 2301600011 generated successfully!")
add_bullet("The Logistics & Printing section becomes active below")
add_bullet("Use 'Copy GP Number' to copy the gate pass number to clipboard")
img_placeholder("Success Dialog — Gate Pass 2301600011 Generated Successfully")
doc.add_paragraph()

# ── STEP 13 ────────────────────────────────────────────────────
add_step_heading(13, "Fill Logistics & Documents Details")
add_body(
    "After gate pass generation, the Logistics & Documents section becomes active. "
    "Fill in the transport and document details."
)
add_bullet("DC Option — select 'DC Not Required' or 'DC Required'")
add_bullet("L.R / Vehicle No. — enter the Lorry Receipt or Vehicle number")
add_bullet("Transport By — select 'Self' (auto-fills company name & GST) or 'Vendor' (manual entry)")
add_bullet("Mode of Transport — By Road / By Rail / By Air")
add_bullet("Transporter Name & Transporter GST — enter manually if Vendor transport")
add_bullet("E-Way Bill No. & E-Way Bill Date — enabled only when DC Required is selected")
add_bullet("DC/DF Notes — additional notes for the delivery challan")
add_bullet("Status — set to OPEN, Awaiting for Return, CLOSED, or CANCELLED")
add_bullet("Comments — add any tracking comments with dates")
img_placeholder("Logistics & Documents — DC Not Required option selected")
img_placeholder("Logistics & Documents — DC Required option (shows E-Way Bill & DC Notes fields)")
doc.add_paragraph()

# ── STEP 14 ────────────────────────────────────────────────────
add_step_heading(14, "Save & Print Gate Pass")
add_body(
    "After filling logistics details, the Store User saves and prints the gate pass."
)
add_bullet("Click SAVE to store the logistics details in the backend")
add_bullet("Click PRINT GATE PASS to generate the Gate Pass PDF")
add_bullet("Click GENERATE DC to generate the Delivery Challan PDF (if DC Required)")
add_bullet("Click CANCEL GATE PASS to cancel an already-generated gate pass (if needed)")
img_placeholder("Action Buttons — PRINT GATE PASS, CANCEL GATE PASS, GENERATE DC, SAVE")
doc.add_paragraph()

# ── STEP 15 ────────────────────────────────────────────────────
add_step_heading(15, "Delivery Challan (DC) — Generated PDF")
add_body(
    "When DC Required is selected and 'Generate DC' is clicked, a Delivery Challan PDF is produced."
)
add_bullet("Company header: MEIL Neyveli Energy Private Limited with GSTIN & CIN")
add_bullet("DC No., DC Date, GP No., GP Date")
add_bullet("Mode of Transport, LR/Vehicle No., Transporter Name")
add_bullet("Despatch From / Despatch To locations")
add_bullet("Item table: Description, HSN Code, UOM, Quantity, Rate, Amount")
img_placeholder("Delivery Challan PDF — DC/2026/2301600011")
doc.add_paragraph()

# ── STEP 16 ────────────────────────────────────────────────────
add_step_heading(16, "Gate Pass Print — Final Printed Document")
add_body(
    "Clicking PRINT GATE PASS generates the official NON-RETURNABLE GATE PASS PDF "
    "which is presented at the plant gate for security clearance."
)
add_bullet("Company header: MEIL Neyveli Energy Private Limited")
add_bullet("GP No. and Date (top right)")
add_bullet("Vendor Name and Address")
add_bullet("Gate Pass Type: NRGP")
add_bullet("Department, Vehicle No., Vendor GST")
add_bullet("Item table: Description, HSN Code, Outward QTY, UOM, Rate, Value")
add_bullet("Total Value in words and figures")
add_bullet("Req. No., Requestor, Department, Approved By, DC No., Mode")
add_note("The printed gate pass is handed to the security at the gate. The security verifies the details and allows the vehicle to exit the premises.")
img_placeholder("Non-Returnable Gate Pass (NRGP) Print PDF — GP No: 2301600011")
doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  AMENDMENT FLOW
# ═══════════════════════════════════════════════════════════════
add_heading("4.  Amendment Flow", 1)
add_body(
    "If the HOD finds an issue with the request, they can click AMENDMENT in SAP Inbox "
    "instead of Approve. This sends the request back to the original Requestor for correction."
)
doc.add_paragraph()

amend_steps = [
    ("HOD sends Amendment",
     "HOD clicks 'AMENDMENT' in SAP My Inbox. The request status changes to 'Amendment' in the Gate Pass Management System."),
    ("Requestor notified",
     "The Requestor opens the Gate Pass Requests List. The request appears with status 'Amendment' (shown in orange). "
     "The Request No. is now a clickable blue link."),
    ("Requestor opens & edits",
     "The Requestor clicks the Request No. link → the Out Gate Pass Generation screen opens in editable mode. "
     "All header fields (Plant, Vendor, Department) and remarks are editable."),
    ("Requestor re-submits",
     "After making corrections, the Requestor clicks the 'RE-SUBMIT REQUEST' button. "
     "The request is re-sent for approval to HOD and Store."),
    ("Approval restarts",
     "HOD and Store User review and approve the corrected request again via SAP Inbox."),
]
for i, (title, desc) in enumerate(amend_steps, 1):
    p = doc.add_paragraph()
    run_no = p.add_run(f"  {i}.  ")
    run_no.bold = True
    run_no.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
    run_no.font.size = Pt(11)
    run_title = p.add_run(title + " — ")
    run_title.bold = True
    run_title.font.size = Pt(11)
    p.add_run(desc).font.size = Pt(11)

doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  STATUS REFERENCE
# ═══════════════════════════════════════════════════════════════
add_heading("5.  Status Reference", 1)
status_data = [
    ("Pending",   "Warning (Orange)", "Request submitted; waiting for HOD or Store approval"),
    ("Approved",  "Success (Green)",  "Approved by both HOD and Store; ready for Gate Pass generation"),
    ("Rejected",  "Error (Red)",      "Rejected by HOD or Store; no further action possible"),
    ("Amendment", "Warning (Orange)", "Sent back by HOD for correction; Requestor must edit & re-submit"),
    ("Closed",    "Success (Green)",  "Gate Pass has been closed after material dispatch"),
    ("Cancelled", "Error (Red)",      "Gate Pass has been cancelled"),
]
tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = 'Table Grid'
h = tbl2.rows[0].cells
h[0].text = "Status"
h[1].text = "Indicator"
h[2].text = "Description"
for cell in h:
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    cell._tc.get_or_add_tcPr().append(OxmlElement('w:shd'))
    shd = cell._tc.tcPr.find(qn('w:shd'))
    shd.set(qn('w:fill'), '1A237E')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'), 'clear')

for status, indicator, desc in status_data:
    row = tbl2.add_row().cells
    row[0].text = status
    row[1].text = indicator
    row[2].text = desc

doc.add_paragraph()

# Footer note
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— End of Document —")
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
run.font.size = Pt(11)

# ── Save ──────────────────────────────────────────────────────
out_path = r"c:\Fiori\zgpms\NRGP_Process_Guide.docx"
doc.save(out_path)
print(f"Document saved: {out_path}")
